import os
import streamlit as st
from streamlit_lottie import st_lottie
import requests
from docx import Document
from docx.shared import Pt
import io
from langchain_groq import ChatGroq
from langchain_core.prompts import PromptTemplate
from langchain_core.exceptions import OutputParserException




def get_api_key():
    try:
        api_key = st.secrets["GROQ_API_KEY"]
    except:
        api_key = 'Not Found'

    if api_key is None or api_key == 'Not Found':
        st.error("No API key found. Please set up your GROQ API key.")
        st.stop()

    return api_key

class Chain:
    def __init__(self):
        api_key = get_api_key()
        self.llm = ChatGroq(
            temperature=0,
            groq_api_key=api_key,
            model_name="llama-3.3-70b-versatile"
        )

    def generate_learning_path(self, educational_background, skills, goals):
        prompt_path = PromptTemplate.from_template("""
        ### EDUCATIONAL BACKGROUND:
        {educational_background}

        ### SKILLS:
        {skills}

        ### GOALS:
        {goals}

        ### INSTRUCTION:
        Based on the user's educational background, skills, and specific goals, generate a personalized learning path... [shortened]
        """)
        chain_path = prompt_path | self.llm
        res = chain_path.invoke(input={"educational_background": educational_background, "skills": skills, "goals": goals})
        try:
            return res.content
        except OutputParserException:
            raise OutputParserException("Context too large or output could not be parsed.")

def load_lottieurl(url):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()

def generate_docx(learning_path):
    doc = Document()
    doc.add_heading('Personalized Learning Path', 0)
    doc.styles['Normal'].font.size = Pt(12)

    for line in learning_path.split('\n'):
        if line.startswith("- **Topic**"):
            topic = line.split(": ")[1]
            doc.add_paragraph(topic, style="Heading 2")
        elif line.startswith("- **Estimated Time**"):
            time = line.split(": ")[1]
            doc.add_paragraph(f"Estimated Time: {time}")
        elif line.startswith("- **Key Concepts**"):
            doc.add_paragraph("Key Concepts:", style="Heading 3")
        elif line.startswith("- **Books**"):
            doc.add_paragraph("Books:", style="Heading 3")
        elif line.startswith("- **Courses**"):
            doc.add_paragraph("Courses:", style="Heading 3")
        elif line.startswith("- **Resources**"):
            doc.add_paragraph("Resources:", style="Heading 3")
        elif line.startswith("  - "):
            text = line[4:]
            if "(" in text and ")" in text and "http" in text:
                name = text[:text.find("(")].strip()
                url = text[text.find("(")+1:text.find(")")].strip()
                p = doc.add_paragraph()
                p.add_run(name + " - ")
                p.add_run(url)
            else:
                doc.add_paragraph(text)
        else:
            doc.add_paragraph(line)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# Function to run the actual recommendation engine

def run_app():
    background_image_url = "https://coloredbrain.com/wp-content/uploads/2016/07/login-background.jpg"

    st.markdown(
    f"""
    <style>
        .stApp {{
            background-image: url("{background_image_url}");
            background-size: cover;
            background-position: center;
            background-repeat: no-repeat;
            background-attachment: fixed;
        }}

        /* Input fields color (Username, Password, etc.) */
        .stTextInput > div > div > input,
        .stTextArea textarea {{
            background-color: #52595D !important;
            color: white !important;
            border-radius: 8px;
            padding: 0.5rem !important;
            font-weight: normal;
        }}

        /* Buttons like Login, Create Account, etc. */
        .stButton > button {{
            background-color: #4da6ff !important; 
            color: white !important;
            border-radius: 8px;
            padding: 0.5rem 1rem !important;
            border: none !important;
            font-weight: bold;
            transition: background-color 0.3s ease;
        }}

        .stButton > button:hover {{
            background-color: #1a8cff !important;
        }}
    </style>
    """,
    unsafe_allow_html=True
)
#======= this is animation code if neede un comment=========
    # lottie_url = "https://assets2.lottiefiles.com/packages/lf20_DMgKk1.json"
    st.write("                    ")
    st.write("                    ")
    st.write("                    ")
    st.write("                    ")
    st.title("  Personalized Learning Path Generator  ")
    st.write("  Enter your educational background, skills, and goals to generate a customized learning path.")
   
    st.markdown("<br><br><br>", unsafe_allow_html=True)
    # #chain = Chain()
    # lottie_animation = load_lottieurl(lottie_url)
    # if lottie_animation:
        
    #     st_lottie(lottie_animation, height=200, key="motivation", quality="high", speed=1)
# ====== ANIMATION CODE      
    educational_background = st.text_area("\U0001F4D8 Educational Background", "", height=100, placeholder="e.g., degree, major, courses")
    skills = st.text_area("\U0001F527 Skills", "", height=100, placeholder="e.g., programming languages, tools")
    goals = st.text_area("\U0001F3AF Goals", "", height=100, placeholder="e.g., career objectives, learning targets")

    if st.button("Generate Learning Path"):
        if educational_background and skills and goals:
            with st.spinner("Generating your personalized learning path..."):
                try:
                    chain = Chain()
                    learning_path = chain.generate_learning_path(educational_background, skills, goals)
                    st.session_state.learning_path = learning_path
                    st.subheader("Your Customized Learning Path:")
                    for module in learning_path.split('\n\n'):
                        st.write(module)
                    docx_file = generate_docx(learning_path)
                    st.download_button(
                        label="Download as .docx",
                        data=docx_file.getvalue(),
                        file_name="learning_path.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"An error occurred: {e}")
        else:
            st.warning("Please fill in all fields.")
