import streamlit as st
from chains import Chain  
from streamlit_lottie import st_lottie
import requests
from docx import Document
import io

# Initialize Chain class
chain = Chain()

# Load Lottie animation
def load_lottieurl(url):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()

# Lottie animation URL
lottie_url = "https://assets9.lottiefiles.com/packages/lf20_tfb3estd.json"
lottie_animation = load_lottieurl(lottie_url)


# Function to generate .docx from the learning path
def generate_docx(learning_path):
    doc = Document()
    doc.add_heading('Personalized Learning Path', 0)

    for course in learning_path.split('\n'):
        doc.add_paragraph(course)

    # Save document to a BytesIO object
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

# Enhanced Custom CSS for full page styling and animation
st.markdown(
    """
    <style>
        /* Full page background color */
        .stApp {
            background-color: #f7ede2;
            animation: backgroundMove 15s infinite alternate;
        }

        /* Keyframes for subtle background motion */
        @keyframes backgroundMove {
            0% { background-position: 0% 0%; }
            100% { background-position: 100% 100%; }
        }

        /* Main content color (text) */
        .css-1d391kg, .css-10trblm, .css-145kmo2, .css-1629p8f, .css-2trqyj, .css-4yfn50 {
            color: #606c38 !important; /* Contrasting pastel green text */
        }

        /* Title Styling */
        h1 {
            color: #283618 !important; /* Darker olive green for title */
            text-align: center;
            font-weight: bold;
            font-size: 2.3em;
        }

        /* Button Styling */
        .stButton>button {
            background-color: #e9c46a; /* Soft pastel yellow for buttons */
            color: #283618; /* Dark olive green text */
            border: none;
            border-radius: 8px;
            padding: 0.5em 1em;
            font-weight: bold;
            transition: background-color 0.3s ease;
        }

        .stButton>button:hover {
            background-color: #f4a261; /* Muted orange hover effect */
            color: white;
        }

        /* Input Box Styling */
        .stTextArea textarea {
            background-color: #faf3e0; /* Very light beige for input areas */
            color: #283618; /* Dark olive green text for inputs */
            border: 2px solid #e9c46a;
            border-radius: 8px;
            font-size: 1rem;
            padding: 10px;
        }

        /* Subheader Styling */
        h2, h3 {
            color: #606c38 !important; /* Pastel green for subheadings */
            font-weight: bold;
            font-size: 1.8em; /* Increase font size for subheadings */
        }

        /* Separator */
        hr {
            border-top: 2px solid #e9c46a;
        }

        /* Specific text size adjustment for st.write */
        .stMarkdown p {
            font-size: 1.1em; /* Increase font size of the paragraph inside st.write */
        }


        /* Label styling for st.text_area */
        .stTextArea>label {
            font-size: 1.5em !important; /* Increase font size of label */
            font-weight: bold;
        }

        /* Styling for the placeholder text */
        .stTextArea textarea::placeholder {
            font-size: 1em;
        }
    </style>
    """, unsafe_allow_html=True
)

# UI layout
st.title("📘 Personalized Learning Path Generator")
st.write("Enter your educational background, skills, and goals to generate a customized learning path.")

# Place Lottie animation in the main content area
if lottie_animation:
    st.markdown('<div class="lottie-container">', unsafe_allow_html=True)
    st_lottie(lottie_animation, height=200, key="motivation")
    st.markdown('</div>', unsafe_allow_html=True)

# User inputs with descriptions and placeholders
educational_background = st.text_area(
    "📘 Educational Background", 
    "",  # No default value, just an empty field
    height=100,
    placeholder="Enter details (e.g., degree, major, relevant courses):"
)
skills = st.text_area(
    "🔧 Skills", 
    "",  # No default value, just an empty field
    height=100,
    placeholder="Enter details (e.g., programming languages, technical skills, tools):"
)
goals = st.text_area(
    "🎯 Goals", 
    "",  # No default value, just an empty field
    height=100,
    placeholder="Enter details (e.g., career objectives, learning goals):"
)

# Generate learning path button
if st.button("Generate Learning Path"):
    if educational_background and skills and goals:
        with st.spinner("Generating your personalized learning path..."):
            try:
                # Call the generate_learning_path method
                learning_path = chain.generate_learning_path(educational_background, skills, goals)
                
                # Store learning path in session state
                st.session_state.learning_path = learning_path
                
                # Display the learning path in a readable format
                st.subheader("Your Customized Learning Path:")
                
                # Timeline-style display
                st.markdown("<hr>", unsafe_allow_html=True)
                for i, module in enumerate(learning_path.split('\n\n'), 1):
                    st.write(f"{module}")
                    st.markdown("<hr>", unsafe_allow_html=True)

                # Generate the .docx file
                docx_file = generate_docx(learning_path)

                # Provide a download link
                st.download_button(
                    label="Download Learning Path as .docx",
                    data=docx_file.getvalue(),
                    file_name="learning_path.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )    
            except Exception as e:
                st.error(f"An error occurred: {e}")
    else:
        st.warning("Please provide educational background, skills, and goals.")
